"""다포맷 텍스트 추출기 — HTML / HWP / PDF / DOCX / RTF.

각 추출기는 path → str 단순 인터페이스.
HWP는 Windows + 한컴오피스 환경에서만 작동.
"""
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional
from .signatures import detect_kind, FileKind


# ───────────────── HTML ─────────────────
def extract_html(path: Path) -> str:
    from bs4 import BeautifulSoup
    raw = Path(path).read_bytes()
    for enc in ('utf-8', 'cp949', 'euc-kr'):
        try:
            txt = raw.decode(enc); break
        except UnicodeDecodeError:
            continue
    else:
        txt = raw.decode('utf-8', errors='ignore')
    soup = BeautifulSoup(txt, 'html.parser')
    for s in soup(['script', 'style', 'noscript']):
        s.decompose()
    return soup.get_text(separator='\n')


# ───────────────── PDF ─────────────────
def extract_pdf(path: Path) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ''
            out.append(t)
    return '\n'.join(out)


# ───────────────── HWP (Windows + 한컴오피스) ─────────────────
def extract_hwp(path: Path, hwp_obj=None) -> str:
    """HwpObject 재사용 가능 (배치 처리 시 매번 생성 비용 큼)."""
    import win32com.client
    import pythoncom

    own_obj = hwp_obj is None
    if own_obj:
        pythoncom.CoInitialize()
        hwp_obj = win32com.client.Dispatch('HWPFrame.HwpObject')
        try:
            hwp_obj.SetMessageBoxMode(0x10)
        except Exception:
            pass

    try:
        ok = hwp_obj.Open(str(Path(path).absolute()), "",
                          "forceopen:true;suspendpassword:true")
        if not ok:
            return '<HWP_OPEN_FAILED>'
        text = hwp_obj.GetTextFile("TEXT", "")
        try:
            hwp_obj.XHwpDocuments.Item(0).Close(isDirty=False)
        except Exception:
            try: hwp_obj.Clear(1)
            except Exception: pass
        return text or ''
    finally:
        if own_obj:
            try: hwp_obj.Quit()
            except Exception: pass
            pythoncom.CoUninitialize()


def extract_hwp_via_temp(path: Path, hwp_obj, suffix='.hwp') -> str:
    """page.html이 실제로는 HWP 바이너리일 때 임시 .hwp 복사 후 처리."""
    import shutil
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(Path(path).read_bytes())
        tmp_path = Path(tmp.name)
    try:
        return extract_hwp(tmp_path, hwp_obj)
    finally:
        try: tmp_path.unlink()
        except Exception: pass


# ───────────────── DOCX ─────────────────
def extract_docx(path: Path) -> str:
    """python-docx 사용. Windows·Mac·Linux 공통."""
    try:
        import docx
    except ImportError:
        return '<DOCX_MODULE_MISSING — pip install python-docx>'
    doc = docx.Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


# ───────────────── RTF ─────────────────
def extract_rtf(path: Path) -> str:
    """striprtf(Python) 사용. unrtf 없어도 동작."""
    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(Path(path).read_text(encoding='utf-8', errors='ignore'))
    except ImportError:
        # fallback to unrtf if installed
        r = subprocess.run(['unrtf', '--text', str(path)],
                          capture_output=True, timeout=60)
        return r.stdout.decode('utf-8', errors='ignore')


# ───────────────── 시그니처 기반 자동 분기 ─────────────────
def extract_by_signature(path: Path, hwp_obj=None) -> tuple:
    """반환: (extracted_text, detected_kind)"""
    kind = detect_kind(path)
    if kind == FileKind.PDF:
        return extract_pdf(path), kind
    if kind in (FileKind.OLE2, FileKind.HWP_ALT):
        if hwp_obj:
            return extract_hwp(path, hwp_obj), kind
        return extract_hwp(path), kind
    if kind == FileKind.ZIP:
        # DOCX/HWPX 시도 — 일단 docx 우선
        try:
            return extract_docx(path), kind
        except Exception as e:
            return f'<ZIP_NOT_HANDLED: {e}>', kind
    if kind == FileKind.HTML:
        return extract_html(path), kind
    if kind == FileKind.RTF:
        return extract_rtf(path), kind
    return f'<UNKNOWN_FORMAT>', kind


def normalize_text(text: str) -> str:
    """공통 후처리.

    수행 단계:
    1. HTML 엔터티 디코딩 (&#9702; → ◇, &amp; → & 등)
    2. 줄바꿈 정리
    3. 전각/특수 글머리표 → 공백 (S5/S2 정확도 향상)
    4. 양식·약어 풀이 잔재 컷
    5. 연속 공백 압축
    """
    import html as html_lib
    # 1. HTML 엔터티 디코딩
    text = html_lib.unescape(text)
    # 2. 줄바꿈 정리
    text = re.sub(r'\r\n?', '\n', text)
    # 3. 전각/특수 글머리표 → 세미콜론 (sub-requirement 분리 + 노이즈 제거)
    #    예: "관리방안을 제시 ◇ 보고체계를 정립 ◇ 일정 제출" → "...제시; 보고체계 정립; 일정 제출"
    text = re.sub(r'\s*[◇□■○●▶▪◐▷◆＊]\s*', '; ', text)
    #    전각 하이픈·대시 (단독 글머리표) → 세미콜론
    text = re.sub(r'(^|\s)[－–—]\s+', r'\1; ', text)
    #    인용·문장부호류 → 공백
    text = re.sub(r'[「」『』※]', ' ', text)
    #    세미콜론 정리 (연속 세미콜론, 시작 세미콜론)
    text = re.sub(r'(;\s*){2,}', '; ', text)
    text = re.sub(r'^\s*;\s*', '', text, flags=re.MULTILINE)
    # 4. 양식·약어 풀이 잔재 컷
    text = re.sub(r'\[양식\s*\d+\s*(?:참고|참조)?\]', '', text)
    text = re.sub(r'\[별첨\s*\d+\]', '', text)
    text = re.sub(r'\[붙임\s*\d+\]', '', text)
    text = re.sub(r'\(M/M\)', '', text)
    # 5. 공백 압축
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    lines = [ln.strip() for ln in text.split('\n')]
    return '\n'.join(ln for ln in lines if ln)
